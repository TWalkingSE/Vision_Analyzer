#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
📤 Export Manager - Sistema de Exportação Multi-formato
=======================================================
Exporta relatórios em MD, PDF, DOCX, JSON e HTML.
"""

import base64
import json
import logging
import re
from io import BytesIO
from pathlib import Path
from datetime import datetime
from typing import Optional, Dict, Any, List
from dataclasses import dataclass, asdict, field

logger = logging.getLogger(__name__)

# Verificar dependências opcionais
try:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logger.warning("⚠️ reportlab não instalado. Exportação PDF desabilitada.")

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("⚠️ python-docx não instalado. Exportação DOCX desabilitada.")


PDF_SIDE_MARGIN_CM = 2.4
PDF_TOP_MARGIN_CM = 2.2
PDF_BOTTOM_MARGIN_CM = 2.2


@dataclass
class ReportData:
    """Dados estruturados de um relatório."""
    image_name: str
    image_path: str
    model: str
    analysis_mode: str
    ocr_engine: str
    timestamp: str
    processing_time: float
    dimensions: tuple
    file_size: int
    ocr_result: str
    analysis_content: str
    hash_md5: str
    hash_sha256: str = ""
    image_extension: str = ""
    yolo_result: str = ""
    quality_result: str = ""
    exif_result: str = ""
    post_processing: Dict[str, Any] = field(default_factory=dict)
    post_processing_markdown: str = ""
    images: list[tuple[str, bytes]] = field(default_factory=list)
    preflight_warnings: List[str] = field(default_factory=list)
    pipeline_telemetry: List[Dict[str, Any]] = field(default_factory=list)


class ExportManager:
    """Gerenciador de exportação de relatórios."""
    
    SUPPORTED_FORMATS = ['md', 'json', 'html']
    
    if PDF_AVAILABLE:
        SUPPORTED_FORMATS.append('pdf')
    if DOCX_AVAILABLE:
        SUPPORTED_FORMATS.append('docx')
    
    def __init__(self, output_dir: Path):
        self.output_dir = output_dir
        self.output_dir.mkdir(parents=True, exist_ok=True)
    
    def export(
        self, 
        data: ReportData, 
        formats: List[str] = None,
        base_filename: str = None
    ) -> Dict[str, Path]:
        """
        Exporta relatório em múltiplos formatos.
        Retorna dict com caminhos dos arquivos gerados.
        """
        if formats is None:
            formats = ['md']
        
        if base_filename is None:
            base_filename = f"{data.image_name}_{data.model}"
        
        results = {}
        
        for fmt in formats:
            fmt = fmt.lower()
            if fmt not in self.SUPPORTED_FORMATS:
                logger.warning(f"⚠️ Formato '{fmt}' não suportado")
                continue
            
            try:
                if fmt == 'md':
                    path = self._export_markdown(data, base_filename)
                elif fmt == 'json':
                    path = self._export_json(data, base_filename)
                elif fmt == 'html':
                    path = self._export_html(data, base_filename)
                elif fmt == 'pdf':
                    path = self._export_pdf(data, base_filename)
                elif fmt == 'docx':
                    path = self._export_docx(data, base_filename)
                else:
                    continue
                
                results[fmt] = path
                logger.info(f"📄 Exportado: {path.name}")
                
            except Exception as e:
                logger.error(f"❌ Erro ao exportar {fmt}: {e}")
        
        return results

    @staticmethod
    def _text_or_placeholder(value: str, placeholder: str = "[Não disponível]") -> str:
        """Retorna um valor textual com fallback consistente."""
        return value if value else placeholder

    def _pre_analysis_markdown(self, data: ReportData) -> str:
        """Monta a seção Markdown de pré-análise."""
        return f"""## 🔍 Dados Extraídos (Pré-Análise)

### EXIF e GPS (Metadados Ocultos)
```text
{self._text_or_placeholder(data.exif_result, '[EXIF não avaliado]')}
```

### OCR (Texto Visível)
```text
{self._text_or_placeholder(data.ocr_result, '[Nenhum texto detectado]')}
```

### YOLO11 (Objetos)
```text
{self._text_or_placeholder(data.yolo_result, '[YOLO não executado]')}
```

### Qualidade da Imagem
```text
{self._text_or_placeholder(data.quality_result, '[Qualidade não avaliada]')}
```
"""

    def _post_processing_markdown(self, data: ReportData) -> str:
        """Monta a seção Markdown de pós-processamento."""
        if data.post_processing_markdown:
            return data.post_processing_markdown.strip()
        return ""

    def _pipeline_markdown(self, data: ReportData) -> str:
        """Monta a seção Markdown com avisos e telemetria da pipeline."""
        if not data.preflight_warnings and not data.pipeline_telemetry:
            return ""

        blocks = ["## ⚙️ Execução da Pipeline\n"]

        if data.preflight_warnings:
            warnings = "\n".join(f"- {warning}" for warning in data.preflight_warnings)
            blocks.append(f"### Avisos de Preflight\n{warnings}\n")

        if data.pipeline_telemetry:
            rows = ["| Etapa | Status | Duração | Detalhe |", "|:------|:-------|--------:|:--------|"]
            for entry in data.pipeline_telemetry:
                detail = str(entry.get("detail", "")).replace("|", "/") or "n/a"
                rows.append(
                    f"| {entry.get('stage', 'n/a')} | {entry.get('status', 'n/a')} | {entry.get('duration_ms', 0)} ms | {detail} |"
                )
            blocks.append("### Telemetria por Etapa\n" + "\n".join(rows) + "\n")

        return "\n".join(blocks).strip()

    def _json_payload(self, data: ReportData) -> Dict[str, Any]:
        """Gera o payload JSON enriquecido do relatório."""
        return {
            "metadata": {
                "generator": "Vision Analyzer",
                "version": "2.0.0",
                "exported_at": datetime.now().isoformat(),
            },
            "image": {
                "name": data.image_name,
                "path": data.image_path,
                "extension": data.image_extension,
                "dimensions": list(data.dimensions),
                "file_size": data.file_size,
                "hash_md5": data.hash_md5,
                "hash_sha256": data.hash_sha256,
            },
            "analysis": {
                "model": data.model,
                "mode": data.analysis_mode,
                "ocr_engine": data.ocr_engine,
                "timestamp": data.timestamp,
                "processing_time": data.processing_time,
                "ocr_result": data.ocr_result,
                "content": data.analysis_content,
                "pre_analysis": {
                    "exif_result": data.exif_result,
                    "yolo_result": data.yolo_result,
                    "quality_result": data.quality_result,
                },
                "pipeline": {
                    "preflight_warnings": data.preflight_warnings,
                    "telemetry": data.pipeline_telemetry,
                },
                "post_processing": data.post_processing,
                "post_processing_markdown": data.post_processing_markdown,
            },
        }

    @staticmethod
    def _to_serializable_report_dict(data: ReportData) -> Dict[str, Any]:
        """Serializa o relatório removendo campos binários não exportáveis."""
        payload = asdict(data)
        payload.pop("images", None)
        return payload

    @staticmethod
    def _inline_png_src(image_bytes: bytes) -> str:
        """Converte bytes PNG em data URL inline para HTML."""
        return f"data:image/png;base64,{base64.b64encode(image_bytes).decode('ascii')}"

    def _report_images_html(self, data: ReportData) -> str:
        """Renderiza a galeria inline para HTML/PDF via HTML."""
        if not data.images:
            return ""

        figures = []
        for description, image_bytes in data.images:
            escaped_description = self._escape_xml(description)
            figures.append(
                f"""
        <figure class="image-figure">
            <img class="report-image" src="{self._inline_png_src(image_bytes)}" alt="{escaped_description}">
            <figcaption>{escaped_description}</figcaption>
        </figure>
"""
            )

        return f"""
    <div class="card">
        <h2>🖼️ Imagens Analisadas</h2>
        <div class="image-gallery">
            {''.join(figures)}
        </div>
    </div>
"""

    def _build_pdf_image(self, image_bytes: bytes):
        """Cria uma imagem ReportLab redimensionada proporcionalmente."""
        image = RLImage(BytesIO(image_bytes))
        max_width = 15 * cm
        max_height = 10 * cm
        scale = min(max_width / image.imageWidth, max_height / image.imageHeight, 1)
        image.drawWidth = image.imageWidth * scale
        image.drawHeight = image.imageHeight * scale
        return image

    def _report_images_pdf(self, data: ReportData, heading_style, body_style) -> List[Any]:
        """Renderiza as imagens anexadas para exportação PDF."""
        if not data.images:
            return []

        elements = [Paragraph("🖼️ Imagens Analisadas", heading_style), Spacer(1, 10)]
        for description, image_bytes in data.images:
            elements.append(Paragraph(self._escape_xml(description), body_style))
            elements.append(Spacer(1, 4))
            elements.append(self._build_pdf_image(image_bytes))
            elements.append(Spacer(1, 12))
        return elements

    def _pre_analysis_html(self, data: ReportData) -> str:
        """Monta a seção HTML de pré-análise."""
        return f"""
    <div class="card">
        <h2>🔍 Pré-Análise</h2>
        <h3>EXIF e GPS</h3>
        <div class="ocr-box">{self._escape_xml(self._text_or_placeholder(data.exif_result, '[EXIF não avaliado]')).replace(chr(10), '<br>')}</div>
        <h3>YOLO11</h3>
        <div class="ocr-box">{self._escape_xml(self._text_or_placeholder(data.yolo_result, '[YOLO não executado]')).replace(chr(10), '<br>')}</div>
        <h3>Qualidade</h3>
        <div class="ocr-box">{self._escape_xml(self._text_or_placeholder(data.quality_result, '[Qualidade não avaliada]')).replace(chr(10), '<br>')}</div>
    </div>
"""

    def _post_processing_html(self, data: ReportData) -> str:
        """Monta a seção HTML de pós-processamento."""
        if not data.post_processing and not data.post_processing_markdown:
            return ""

        content = data.post_processing_markdown or json.dumps(data.post_processing, indent=2, ensure_ascii=False)
        return f"""
    <div class="card">
        <h2>🧠 Pós-Processamento Estruturado</h2>
        <div class="analysis-content">{self._escape_xml(content).replace(chr(10), '<br>')}</div>
    </div>
"""

    def _pipeline_html(self, data: ReportData) -> str:
        """Monta a seção HTML com avisos e telemetria da pipeline."""
        if not data.preflight_warnings and not data.pipeline_telemetry:
            return ""

        warnings_html = ""
        if data.preflight_warnings:
            warning_items = "".join(f"<li>{self._escape_xml(warning)}</li>" for warning in data.preflight_warnings)
            warnings_html = f"<h3>Avisos de Preflight</h3><ul>{warning_items}</ul>"

        telemetry_html = ""
        if data.pipeline_telemetry:
            rows = []
            for entry in data.pipeline_telemetry:
                rows.append(
                    "<tr>"
                    f"<td>{self._escape_xml(str(entry.get('stage', 'n/a')))}</td>"
                    f"<td>{self._escape_xml(str(entry.get('status', 'n/a')))}</td>"
                    f"<td>{self._escape_xml(str(entry.get('duration_ms', 0)))} ms</td>"
                    f"<td>{self._escape_xml(str(entry.get('detail', '')) or 'n/a')}</td>"
                    "</tr>"
                )
            telemetry_html = (
                "<h3>Telemetria por Etapa</h3>"
                "<table><tr><th>Etapa</th><th>Status</th><th>Duração</th><th>Detalhe</th></tr>"
                + "".join(rows)
                + "</table>"
            )

        return f"""
    <div class="card">
        <h2>⚙️ Execução da Pipeline</h2>
        {warnings_html}
        {telemetry_html}
    </div>
"""
    
    def _export_markdown(self, data: ReportData, base_filename: str) -> Path:
        """Exporta para Markdown."""
        filepath = self.output_dir / f"{base_filename}.md"
        pre_analysis_section = self._pre_analysis_markdown(data)
        pipeline_section = self._pipeline_markdown(data)
        post_processing_section = self._post_processing_markdown(data)
        post_processing_block = f"\n---\n\n{post_processing_section}\n" if post_processing_section else ""
        pipeline_block = f"\n---\n\n{pipeline_section}\n" if pipeline_section else ""
        
        content = f"""---
# 📊 Relatório de Análise de Imagem
---

**Arquivo:** `{data.image_name}`  
**Modelo:** `{data.model}`  
**Modo:** `{data.analysis_mode}`  
**Data:** {data.timestamp}  
**Tempo de Processamento:** {data.processing_time:.2f}s  

---

## 📋 Metadados da Imagem

| Propriedade | Valor |
|:------------|:------|
| **Dimensões** | {data.dimensions[0]} x {data.dimensions[1]} px |
| **Formato Original** | {data.image_extension.upper() if data.image_extension else 'N/D'} |
| **Tamanho** | {data.file_size / 1024:.1f} KB |
| **Hash MD5** | `{data.hash_md5}` |
| **Hash SHA-256** | `{data.hash_sha256}` |
| **OCR Engine** | {data.ocr_engine} |

---

{pre_analysis_section}

---

## 🤖 Análise do Modelo

{self._text_or_placeholder(data.analysis_content, '[Análise não disponível — o modelo não retornou conteúdo válido]')}

{pipeline_block}

{post_processing_block}

---

*Relatório gerado automaticamente por Vision Analyzer*
"""
        filepath.write_text(content, encoding='utf-8')
        return filepath
    
    def _export_json(self, data: ReportData, base_filename: str) -> Path:
        """Exporta para JSON."""
        filepath = self.output_dir / f"{base_filename}.json"
        json_data = self._json_payload(data)
        
        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(json_data, f, indent=2, ensure_ascii=False)
        
        return filepath
    
    def _export_html(self, data: ReportData, base_filename: str) -> Path:
        """Exporta para HTML."""
        filepath = self.output_dir / f"{base_filename}.html"
        pre_analysis_html = self._pre_analysis_html(data)
        pipeline_html = self._pipeline_html(data)
        post_processing_html = self._post_processing_html(data)
        report_images_html = self._report_images_html(data)
        
        # Tema baseado no modo de análise
        if data.analysis_mode == "forense":
            theme_color = "#dc3545"
            theme_bg = "#1a1a2e"
            theme_text = "#e0e0e0"
        else:
            theme_color = "#007bff"
            theme_bg = "#f8f9fa"
            theme_text = "#212529"
        
        html_content = f"""<!DOCTYPE html>
<html lang="pt-BR">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Relatório - {data.image_name}</title>
    <style>
        :root {{
            --primary: {theme_color};
            --bg: {theme_bg};
            --text: {theme_text};
        }}
        body {{
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            background: var(--bg);
            color: var(--text);
            max-width: 900px;
            margin: 0 auto;
            padding: 20px;
            line-height: 1.6;
        }}
        .header {{
            background: linear-gradient(135deg, var(--primary), #6c757d);
            color: white;
            padding: 30px;
            border-radius: 10px;
            margin-bottom: 20px;
        }}
        .header h1 {{
            margin: 0;
            font-size: 1.8em;
        }}
        .header .subtitle {{
            opacity: 0.9;
            margin-top: 10px;
        }}
        .card {{
            background: white;
            color: #1f1f1f;
            border-radius: 10px;
            padding: 20px;
            margin-bottom: 20px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        }}
        .card h2 {{
            color: var(--primary);
            border-bottom: 2px solid var(--primary);
            padding-bottom: 10px;
        }}
        table {{
            width: 100%;
            border-collapse: collapse;
        }}
        th, td {{
            padding: 10px;
            text-align: left;
            border-bottom: 1px solid #ddd;
        }}
        th {{
            background: var(--primary);
            color: white;
        }}
        .ocr-box {{
            background: #f5f5f5;
            border-left: 4px solid var(--primary);
            padding: 15px;
            font-family: monospace;
            white-space: pre-wrap;
        }}
        .analysis-content {{
            background: #fafafa;
            color: inherit;
            padding: 20px;
            border-radius: 5px;
        }}
        .footer {{
            text-align: center;
            color: #888;
            margin-top: 30px;
            padding: 20px;
        }}
        .badge {{
            display: inline-block;
            padding: 5px 10px;
            border-radius: 15px;
            background: var(--primary);
            color: white;
            font-size: 0.8em;
            margin-right: 5px;
        }}
        .image-gallery {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(240px, 1fr));
            gap: 16px;
        }}
        .image-figure {{
            margin: 0;
        }}
        .report-image {{
            width: 100%;
            height: auto;
            border-radius: 8px;
            border: 1px solid #d9d9d9;
            background: #fff;
        }}
        figcaption {{
            font-size: 0.9em;
            color: #666;
            margin-top: 8px;
        }}
    </style>
</head>
<body>
    <div class="header">
        <h1>📊 Relatório de Análise de Imagem</h1>
        <div class="subtitle">
            <span class="badge">{data.model}</span>
            <span class="badge">{data.analysis_mode.upper()}</span>
            <span class="badge">{data.ocr_engine}</span>
        </div>
    </div>
    
    <div class="card">
        <h2>📋 Metadados</h2>
        <table>
            <tr><th>Propriedade</th><th>Valor</th></tr>
            <tr><td>Arquivo</td><td><code>{data.image_name}</code></td></tr>
            <tr><td>Dimensões</td><td>{data.dimensions[0]} x {data.dimensions[1]} px</td></tr>
            <tr><td>Formato</td><td>{data.image_extension.upper() if data.image_extension else 'N/D'}</td></tr>
            <tr><td>Tamanho</td><td>{data.file_size / 1024:.1f} KB</td></tr>
            <tr><td>Hash MD5</td><td><code>{data.hash_md5}</code></td></tr>
            <tr><td>Hash SHA-256</td><td><code style="font-size:0.75em">{data.hash_sha256}</code></td></tr>
            <tr><td>Data</td><td>{data.timestamp}</td></tr>
            <tr><td>Tempo</td><td>{data.processing_time:.2f}s</td></tr>
        </table>
    </div>

    {report_images_html}
    
    <div class="card">
        <h2>🔍 Texto Extraído (OCR)</h2>
        <div class="ocr-box">{self._escape_xml(self._text_or_placeholder(data.ocr_result, '[Nenhum texto detectado]')).replace(chr(10), '<br>')}</div>
    </div>

    {pre_analysis_html}

    {pipeline_html}
    
    <div class="card">
        <h2>🤖 Análise do Modelo</h2>
        <div class="analysis-content">
            {self._escape_xml(self._text_or_placeholder(data.analysis_content, '[Análise não disponível — o modelo não retornou conteúdo válido]')).replace(chr(10), '<br>')}
        </div>
    </div>

    {post_processing_html}
    
    <div class="footer">
        <p>Gerado por <strong>Vision Analyzer</strong> | {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
    </div>
</body>
</html>"""
        
        filepath.write_text(html_content, encoding='utf-8')
        return filepath
    
    def _export_pdf(self, data: ReportData, base_filename: str) -> Path:
        """Exporta para PDF com renderização de Markdown."""
        if not PDF_AVAILABLE:
            raise RuntimeError("reportlab não instalado")
        
        filepath = self.output_dir / f"{base_filename}.pdf"
        side_margin = PDF_SIDE_MARGIN_CM * cm
        top_margin = PDF_TOP_MARGIN_CM * cm
        bottom_margin = PDF_BOTTOM_MARGIN_CM * cm
        
        doc = SimpleDocTemplate(
            str(filepath),
            pagesize=A4,
            rightMargin=side_margin,
            leftMargin=side_margin,
            topMargin=top_margin,
            bottomMargin=bottom_margin
        )
        
        styles = getSampleStyleSheet()
        metadata_label_width = 4.2 * cm
        metadata_value_width = doc.width - metadata_label_width
        
        # Estilos customizados
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=20,
            textColor=colors.HexColor('#007bff'),
            alignment=TA_CENTER
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceBefore=15,
            spaceAfter=10,
            textColor=colors.HexColor('#333')
        )
        
        h3_style = ParagraphStyle(
            'CustomH3',
            parent=styles['Heading3'],
            fontSize=12,
            spaceBefore=10,
            spaceAfter=6,
            textColor=colors.HexColor('#555')
        )
        
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontSize=10,
            spaceAfter=8,
            leading=14,
            alignment=TA_JUSTIFY
        )

        body_left_style = ParagraphStyle(
            'CustomBodyLeft',
            parent=body_style,
            alignment=TA_LEFT
        )

        code_style = ParagraphStyle(
            'CustomCode',
            parent=styles['Normal'],
            fontName='Courier',
            fontSize=9,
            leading=12,
            leftIndent=10,
            rightIndent=10,
            spaceBefore=4,
            spaceAfter=8,
            alignment=TA_LEFT,
            backColor=colors.HexColor('#f5f5f5')
        )
        
        bullet_style = ParagraphStyle(
            'CustomBullet',
            parent=styles['Normal'],
            fontSize=10,
            spaceAfter=4,
            leftIndent=20,
            leading=14,
            alignment=TA_LEFT
        )
        
        elements = []
        
        # Título
        elements.append(Paragraph("📊 Relatório de Análise de Imagem", title_style))
        elements.append(Spacer(1, 10))
        
        # Info básica
        elements.append(Paragraph(f"<b>Arquivo:</b> {data.image_name}", body_style))
        elements.append(Paragraph(f"<b>Modelo:</b> {data.model}", body_style))
        elements.append(Paragraph(f"<b>Modo:</b> {data.analysis_mode}", body_style))
        elements.append(Paragraph(f"<b>Data:</b> {data.timestamp}", body_style))
        elements.append(Spacer(1, 15))
        
        # Tabela de metadados
        elements.append(Paragraph("📋 Metadados da Imagem", heading_style))
        
        table_data = [
            ['Propriedade', 'Valor'],
            ['Dimensões', f'{data.dimensions[0]} x {data.dimensions[1]} px'],
            ['Tamanho', f'{data.file_size / 1024:.1f} KB'],
            ['Hash MD5', data.hash_md5[:20] + '...'],
            ['SHA-256', data.hash_sha256[:20] + '...' if data.hash_sha256 else ''],
            ['Tempo', f'{data.processing_time:.2f}s']
        ]
        
        table = Table(table_data, colWidths=[metadata_label_width, metadata_value_width])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#007bff')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#f8f9fa')),
            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#dee2e6'))
        ]))
        elements.append(table)
        elements.append(Spacer(1, 15))

        elements.extend(self._report_images_pdf(data, heading_style, body_left_style))
        
        # OCR
        elements.append(Paragraph("🔍 Texto Extraído (OCR)", heading_style))
        ocr_text = data.ocr_result if data.ocr_result else "[Nenhum texto detectado]"
        elements.append(Paragraph(self._escape_xml(ocr_text).replace('\n', '<br/>'), body_left_style))
        elements.append(Spacer(1, 15))
        
        # Análise — renderização Markdown
        elements.append(Paragraph("🤖 Análise do Modelo", heading_style))
        elements.extend(self._markdown_to_pdf_elements(
            data.analysis_content, body_style, heading_style, h3_style, bullet_style, code_style
        ))
        
        # Footer
        elements.append(Spacer(1, 30))
        footer_style = ParagraphStyle('Footer', parent=styles['Normal'], fontSize=8, textColor=colors.gray, alignment=TA_CENTER)
        elements.append(Paragraph(f"Gerado por Vision Analyzer | {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", footer_style))
        
        doc.build(elements)
        return filepath
    
    @staticmethod
    def _escape_xml(text: str) -> str:
        """Escapa caracteres especiais para XML/ReportLab."""
        text = text.replace('&', '&amp;')
        text = text.replace('<', '&lt;')
        text = text.replace('>', '&gt;')
        text = text.replace('"', '&quot;')
        text = text.replace("'", '&#39;')
        return text
    
    def _flush_pdf_paragraph(self, elements: List[Any], paragraph_lines: List[str], body_style) -> None:
        """Converte linhas corridas em um único parágrafo para justificação adequada."""
        if not paragraph_lines:
            return

        text = " ".join(line.strip() for line in paragraph_lines if line.strip())
        text = re.sub(r'\s+', ' ', text).strip()
        if text:
            elements.append(Paragraph(self._md_inline_to_rl(text), body_style))
        paragraph_lines.clear()

    def _markdown_to_pdf_elements(self, md_text: str, body_style, h2_style, h3_style, bullet_style, code_style=None):
        """Converte texto Markdown em elementos ReportLab."""
        elements = []
        lines = md_text.split('\n')
        
        i = 0
        table_rows = []
        in_table = False
        in_code_block = False
        code_lines: List[str] = []
        paragraph_lines: List[str] = []
        
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            if stripped.startswith('```'):
                if in_table and table_rows:
                    self._flush_pdf_paragraph(elements, paragraph_lines, body_style)
                    elements.append(self._build_md_table(table_rows))
                    table_rows = []
                    in_table = False

                if in_code_block:
                    code_text = self._escape_xml("\n".join(code_lines)).replace('\n', '<br/>')
                    elements.append(Paragraph(code_text, code_style or body_style))
                    code_lines = []
                    in_code_block = False
                else:
                    self._flush_pdf_paragraph(elements, paragraph_lines, body_style)
                    in_code_block = True
                i += 1
                continue

            if in_code_block:
                code_lines.append(line)
                i += 1
                continue
            
            # Linha vazia
            if not stripped:
                if in_table and table_rows:
                    self._flush_pdf_paragraph(elements, paragraph_lines, body_style)
                    elements.append(self._build_md_table(table_rows))
                    table_rows = []
                    in_table = False
                else:
                    self._flush_pdf_paragraph(elements, paragraph_lines, body_style)
                i += 1
                continue
            
            # Tabela Markdown (linha com |)
            if '|' in stripped and stripped.startswith('|'):
                self._flush_pdf_paragraph(elements, paragraph_lines, body_style)
                cells = [c.strip() for c in stripped.strip('|').split('|')]
                # Pular linhas separadoras (|---|---|)
                if all(re.match(r'^:?-+:?$', c) for c in cells if c):
                    in_table = True
                    i += 1
                    continue
                table_rows.append(cells)
                in_table = True
                i += 1
                continue
            
            # Se estávamos em tabela e agora não, fechar
            if in_table and table_rows:
                elements.append(self._build_md_table(table_rows))
                table_rows = []
                in_table = False
            
            # Heading ## 
            if stripped.startswith('## '):
                self._flush_pdf_paragraph(elements, paragraph_lines, body_style)
                text = self._md_inline_to_rl(stripped[3:])
                elements.append(Paragraph(text, h2_style))
                i += 1
                continue
            
            # Heading ###
            if stripped.startswith('### '):
                self._flush_pdf_paragraph(elements, paragraph_lines, body_style)
                text = self._md_inline_to_rl(stripped[4:])
                elements.append(Paragraph(text, h3_style))
                i += 1
                continue
            
            # Bullet point (* ou -)
            if stripped.startswith('* ') or stripped.startswith('- '):
                self._flush_pdf_paragraph(elements, paragraph_lines, body_style)
                text = self._md_inline_to_rl(stripped[2:])
                elements.append(Paragraph(f"• {text}", bullet_style))
                i += 1
                continue
            
            paragraph_lines.append(stripped)
            i += 1

        if in_code_block and code_lines:
            code_text = self._escape_xml("\n".join(code_lines)).replace('\n', '<br/>')
            elements.append(Paragraph(code_text, code_style or body_style))

        self._flush_pdf_paragraph(elements, paragraph_lines, body_style)
        
        # Fechar tabela pendente
        if in_table and table_rows:
            elements.append(self._build_md_table(table_rows))
        
        return elements
    
    def _md_inline_to_rl(self, text: str) -> str:
        """Converte formatação inline Markdown para tags ReportLab."""
        text = self._escape_xml(text)
        # **bold**
        text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
        # *italic*
        text = re.sub(r'\*(.+?)\*', r'<i>\1</i>', text)
        # `code`
        text = re.sub(r'`(.+?)`', r'<font face="Courier">\1</font>', text)
        return text
    
    def _build_md_table(self, rows: list) -> Table:
        """Constrói uma tabela ReportLab a partir de linhas Markdown."""
        if not rows:
            return Spacer(1, 0)
        
        n_cols = max(len(r) for r in rows)
        # Normalizar colunas
        normalized = []
        for r in rows:
            row = list(r) + [''] * (n_cols - len(r))
            normalized.append(row)
        
        col_width = 14 * cm / n_cols
        table = Table(normalized, colWidths=[col_width] * n_cols)
        
        style_cmds = [
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#dee2e6')),
        ]
        
        if len(normalized) > 1:
            style_cmds.extend([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#007bff')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#f8f9fa')),
            ])
        
        table.setStyle(TableStyle(style_cmds))
        return table
    
    def _export_docx(self, data: ReportData, base_filename: str) -> Path:
        """Exporta para DOCX (Word)."""
        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx não instalado")
        
        filepath = self.output_dir / f"{base_filename}.docx"
        
        doc = Document()
        
        # Título
        title = doc.add_heading('Relatório de Análise de Imagem', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Info básica
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run('Arquivo: ').bold = True
        p.add_run(data.image_name)
        
        p = doc.add_paragraph()
        p.add_run('Modelo: ').bold = True
        p.add_run(data.model)
        
        p = doc.add_paragraph()
        p.add_run('Modo: ').bold = True
        p.add_run(data.analysis_mode.upper())
        
        p = doc.add_paragraph()
        p.add_run('Data: ').bold = True
        p.add_run(data.timestamp)
        
        p = doc.add_paragraph()
        p.add_run('Tempo de Processamento: ').bold = True
        p.add_run(f'{data.processing_time:.2f}s')
        
        # Metadados
        doc.add_heading('Metadados da Imagem', level=1)
        
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'
        
        cells = table.rows[0].cells
        cells[0].text = 'Propriedade'
        cells[1].text = 'Valor'
        
        metadata = [
            ('Dimensões', f'{data.dimensions[0]} x {data.dimensions[1]} px'),
            ('Tamanho', f'{data.file_size / 1024:.1f} KB'),
            ('Hash MD5', data.hash_md5),
            ('Hash SHA-256', data.hash_sha256),
            ('OCR Engine', data.ocr_engine)
        ]
        
        for i, (prop, val) in enumerate(metadata, 1):
            cells = table.rows[i].cells
            cells[0].text = prop
            cells[1].text = val
        
        # OCR
        doc.add_heading('Texto Extraído (OCR)', level=1)
        doc.add_paragraph(data.ocr_result)

        # Pré-análise
        doc.add_heading('Pré-Análise', level=1)
        doc.add_paragraph(f"EXIF e GPS\n{self._text_or_placeholder(data.exif_result, '[EXIF não avaliado]')}")
        doc.add_paragraph(f"YOLO11\n{self._text_or_placeholder(data.yolo_result, '[YOLO não executado]')}")
        doc.add_paragraph(f"Qualidade\n{self._text_or_placeholder(data.quality_result, '[Qualidade não avaliada]')}")
        
        # Análise
        doc.add_heading('Análise do Modelo', level=1)
        doc.add_paragraph(data.analysis_content)

        if data.post_processing_markdown:
            doc.add_heading('Pós-Processamento Estruturado', level=1)
            doc.add_paragraph(data.post_processing_markdown)
        
        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run(f'Gerado por Batch Image Analyzer | {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        run.font.size = Pt(8)
        
        doc.save(filepath)
        return filepath
    
    def export_consolidated(
        self, 
        reports: List[ReportData], 
        filename: str = "relatorio_consolidado"
    ) -> Dict[str, Path]:
        """
        Exporta um relatório consolidado com múltiplas análises.
        """
        results = {}
        
        # JSON consolidado
        json_path = self.output_dir / f"{filename}.json"
        json_data = {
            "metadata": {
                "generator": "Vision Analyzer",
                "version": "2.0.0",
                "exported_at": datetime.now().isoformat(),
                "total_reports": len(reports)
            },
            "reports": [self._to_serializable_report_dict(r) for r in reports]
        }
        
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(json_data, f, indent=2, ensure_ascii=False)
        results['json'] = json_path
        
        # HTML consolidado
        html_path = self._export_consolidated_html(reports, filename)
        results['html'] = html_path
        
        return results
    
    def _export_consolidated_html(self, reports: List[ReportData], filename: str) -> Path:
        """Gera HTML consolidado com índice navegável."""
        filepath = self.output_dir / f"{filename}.html"
        
        # Gerar índice
        index_html = ""
        reports_html = ""
        
        for i, data in enumerate(reports):
            anchor = f"report_{i}"
            index_html += f'<li><a href="#{anchor}">{data.image_name} ({data.model})</a></li>\n'
            
            reports_html += f"""
            <div id="{anchor}" class="report-card">
                <h2>📊 {data.image_name}</h2>
                <div class="badges">
                    <span class="badge">{data.model}</span>
                    <span class="badge">{data.analysis_mode}</span>
                </div>
                {self._report_images_html(data) if data.images else ''}
                <p><strong>Tempo:</strong> {data.processing_time:.2f}s</p>
                <details>
                    <summary>🔍 OCR ({len(data.ocr_result)} chars)</summary>
                    <pre>{data.ocr_result[:500]}...</pre>
                </details>
                <details open>
                    <summary>🤖 Análise</summary>
                    <div class="analysis">{data.analysis_content.replace(chr(10), '<br>')}</div>
                </details>
            </div>
            """
        
        html_content = f"""<!DOCTYPE html>
<html lang="pt-BR">
<head>
    <meta charset="UTF-8">
    <title>Relatório Consolidado - {len(reports)} análises</title>
    <style>
        body {{ font-family: 'Segoe UI', sans-serif; max-width: 1200px; margin: 0 auto; padding: 20px; background: #f0f0f0; }}
        .header {{ background: linear-gradient(135deg, #007bff, #6c757d); color: white; padding: 30px; border-radius: 10px; margin-bottom: 20px; }}
        .index {{ background: white; padding: 20px; border-radius: 10px; margin-bottom: 20px; }}
        .index ul {{ columns: 2; }}
        .report-card {{ background: white; padding: 20px; border-radius: 10px; margin-bottom: 15px; }}
        .badge {{ display: inline-block; padding: 3px 10px; background: #007bff; color: white; border-radius: 10px; font-size: 0.8em; margin-right: 5px; }}
        details {{ margin: 10px 0; }}
        summary {{ cursor: pointer; font-weight: bold; }}
        pre {{ background: #f5f5f5; padding: 10px; overflow-x: auto; }}
        .analysis {{ background: #fafafa; padding: 15px; border-radius: 5px; }}
    </style>
</head>
<body>
    <div class="header">
        <h1>📋 Relatório Consolidado</h1>
        <p>{len(reports)} análises | {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
    </div>
    
    <div class="index">
        <h2>📑 Índice</h2>
        <ul>{index_html}</ul>
    </div>
    
    {reports_html}
</body>
</html>"""
        
        filepath.write_text(html_content, encoding='utf-8')
        return filepath


def get_available_formats() -> List[str]:
    """Retorna lista de formatos disponíveis."""
    return ExportManager.SUPPORTED_FORMATS.copy()


def generate_consolidated_pdf(
    reports: List[ReportData],
    output_path: Path,
    title: str = "Relatório Consolidado de Análise",
    include_images: bool = True
) -> Path:
    """
    Gera um PDF consolidado com todas as análises.
    
    Args:
        reports: Lista de ReportData
        output_path: Caminho do arquivo PDF
        title: Título do relatório
        include_images: Se deve incluir thumbnails das imagens
        
    Returns:
        Path do arquivo gerado
    """
    if not PDF_AVAILABLE:
        raise RuntimeError("reportlab não instalado")
    
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        PageBreak, Image as RLImage, KeepTogether
    )
    
    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        rightMargin=PDF_SIDE_MARGIN_CM * cm,
        leftMargin=PDF_SIDE_MARGIN_CM * cm,
        topMargin=PDF_TOP_MARGIN_CM * cm,
        bottomMargin=PDF_BOTTOM_MARGIN_CM * cm
    )
    
    styles = getSampleStyleSheet()
    exporter = ExportManager(output_path.parent)
    meta_label_width = 3.4 * cm
    meta_value_width = doc.width - meta_label_width
    
    # Estilos customizados
    title_style = ParagraphStyle(
        'Title',
        parent=styles['Heading1'],
        fontSize=22,
        spaceAfter=30,
        textColor=colors.HexColor('#1a1a2e'),
        alignment=TA_CENTER
    )
    
    heading_style = ParagraphStyle(
        'Heading',
        parent=styles['Heading2'],
        fontSize=14,
        spaceBefore=20,
        spaceAfter=10,
        textColor=colors.HexColor('#007bff')
    )
    
    subheading_style = ParagraphStyle(
        'SubHeading',
        parent=styles['Heading3'],
        fontSize=12,
        spaceBefore=15,
        spaceAfter=8,
        textColor=colors.HexColor('#333')
    )
    
    body_style = ParagraphStyle(
        'Body',
        parent=styles['Normal'],
        fontSize=9,
        spaceAfter=8,
        leading=12,
        alignment=TA_JUSTIFY
    )

    body_left_style = ParagraphStyle(
        'BodyLeft',
        parent=body_style,
        alignment=TA_LEFT
    )

    bullet_style = ParagraphStyle(
        'Bullet',
        parent=styles['Normal'],
        fontSize=9,
        spaceAfter=4,
        leading=12,
        leftIndent=18,
        alignment=TA_LEFT
    )

    code_style = ParagraphStyle(
        'Code',
        parent=styles['Normal'],
        fontName='Courier',
        fontSize=8,
        leading=11,
        leftIndent=10,
        rightIndent=10,
        spaceBefore=4,
        spaceAfter=8,
        alignment=TA_LEFT,
        backColor=colors.HexColor('#f5f5f5')
    )
    
    small_style = ParagraphStyle(
        'Small',
        parent=styles['Normal'],
        fontSize=8,
        textColor=colors.gray
    )
    
    elements = []
    
    # ========== CAPA ==========
    elements.append(Spacer(1, 3*cm))
    elements.append(Paragraph(f"👁️ {title}", title_style))
    elements.append(Spacer(1, 1*cm))
    
    # Info do relatório
    info_data = [
        ["Data de Geração", datetime.now().strftime('%d/%m/%Y às %H:%M')],
        ["Total de Imagens", str(len(set(r.image_name for r in reports)))],
        ["Total de Análises", str(len(reports))],
        ["Modelos Utilizados", ", ".join(set(r.model for r in reports))],
    ]
    
    info_table = Table(info_data, colWidths=[5 * cm, doc.width - (5 * cm)])
    info_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f0f0f0')),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('PADDING', (0, 0), (-1, -1), 8),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#dee2e6')),
    ]))
    elements.append(info_table)
    
    elements.append(Spacer(1, 2*cm))
    elements.append(Paragraph("Gerado automaticamente por Vision Analyzer v2.0", small_style))
    elements.append(PageBreak())
    
    # ========== ÍNDICE ==========
    elements.append(Paragraph("📑 Índice", heading_style))
    elements.append(Spacer(1, 0.5*cm))
    
    # Agrupar por imagem
    images_dict = {}
    for r in reports:
        if r.image_name not in images_dict:
            images_dict[r.image_name] = []
        images_dict[r.image_name].append(r)
    
    idx = 1
    for img_name, img_reports in images_dict.items():
        elements.append(Paragraph(f"{idx}. {img_name}", body_style))
        for r in img_reports:
            elements.append(Paragraph(f"    • {r.model}", small_style))
        idx += 1
    
    elements.append(PageBreak())
    
    # ========== ANÁLISES ==========
    for img_idx, (img_name, img_reports) in enumerate(images_dict.items(), 1):
        # Cabeçalho da imagem
        elements.append(Paragraph(f"📷 {img_idx}. {img_name}", heading_style))
        
        # Metadados da primeira análise
        first = img_reports[0]
        meta_data = [
            ["Dimensões", f"{first.dimensions[0]} x {first.dimensions[1]} px"],
            ["Tamanho", f"{first.file_size / 1024:.1f} KB"],
            ["Hash MD5", first.hash_md5[:32] + "..."],
            ["SHA-256", first.hash_sha256[:32] + "..." if first.hash_sha256 else ""],
        ]
        
        meta_table = Table(meta_data, colWidths=[meta_label_width, meta_value_width])
        meta_table.setStyle(TableStyle([
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('PADDING', (0, 0), (-1, -1), 4),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ]))
        elements.append(meta_table)
        elements.append(Spacer(1, 0.3*cm))

        if include_images and first.images:
            gallery_exporter = ExportManager(output_path.parent)
            elements.extend(gallery_exporter._report_images_pdf(first, subheading_style, small_style))
        
        # OCR (uma vez por imagem)
        if first.ocr_result and first.ocr_result != "[OCR desabilitado]":
            elements.append(Paragraph("🔍 Texto Extraído (OCR)", subheading_style))
            ocr_text = first.ocr_result[:800]
            if len(first.ocr_result) > 800:
                ocr_text += "..."
            elements.append(Paragraph(exporter._escape_xml(ocr_text).replace('\n', '<br/>'), body_left_style))
            elements.append(Spacer(1, 0.3*cm))
        
        # Análises de cada modelo
        for r in img_reports:
            elements.append(Paragraph(f"🤖 Análise: {r.model}", subheading_style))
            elements.append(Paragraph(f"<i>Modo: {r.analysis_mode} | Tempo: {r.processing_time:.2f}s</i>", small_style))
            elements.append(Spacer(1, 0.2*cm))
            
            # Conteúdo da análise (limitado)
            analysis = r.analysis_content[:3000]
            if len(r.analysis_content) > 3000:
                analysis += "\n\n[... conteúdo truncado ...]"

            elements.extend(
                exporter._markdown_to_pdf_elements(
                    analysis,
                    body_style,
                    subheading_style,
                    small_style,
                    bullet_style,
                    code_style,
                )
            )
            elements.append(Spacer(1, 0.5*cm))
        
        # Quebra de página entre imagens (exceto a última)
        if img_idx < len(images_dict):
            elements.append(PageBreak())
    
    # ========== RODAPÉ FINAL ==========
    elements.append(Spacer(1, 1*cm))
    elements.append(Paragraph("—" * 50, small_style))
    elements.append(Paragraph(
        f"Relatório gerado em {datetime.now().strftime('%d/%m/%Y às %H:%M:%S')} por Vision Analyzer v2.0",
        small_style
    ))
    
    # Construir PDF
    doc.build(elements)
    
    return output_path
